"""Word (.docx) report generator for FAA Part 23 certification loads analysis.

Auto-generates a comprehensive certification report including:
- Aircraft description (geometry, structural model, aero panels, control surfaces)
- Analysis scope and certification standards
- Theory & procedures for load generation/selection (with equations)
- Load case generation methodology
- Critical design load selection and results
- VMT envelope, potato plots, V-n diagrams
- Full load case listing in appendix

Dependencies
------------
python-docx >= 1.0

Usage
-----
>>> from nastaero.loads_analysis.certification.report_docx import DocxReportWriter
>>> writer = DocxReportWriter(config, vn_diagram, matrix, batch_result,
...                           envelope_proc, report, model=model,
...                           plot_dir="cert_results_20260304/")
>>> writer.generate("KC100_Cert_Report.docx")
"""
from __future__ import annotations

import math
import os
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional

import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .aircraft_config import AircraftConfig, SpeedSchedule
from .vn_diagram import VnDiagram
from .load_case_matrix import LoadCaseMatrix
from .batch_runner import BatchResult
from .envelope import EnvelopeProcessor
from .report import CertificationReport, FAR_SECTIONS


# ──────────────────────────────────────────────────────────────────
# Styling helpers
# ──────────────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _format_header_row(row, color_hex: str = "2C3E50"):
    """Format a table header row with dark background and white text."""
    for cell in row.cells:
        _set_cell_shading(cell, color_hex)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)


def _set_table_style(table):
    """Set consistent table styling."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Reduce cell padding for compact tables
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)


def _add_alt_row_shading(table, start_row: int = 1,
                          color_hex: str = "EBF5FB"):
    """Add alternating row shading for readability."""
    for i, row in enumerate(table.rows):
        if i < start_row:
            continue
        if (i - start_row) % 2 == 1:
            for cell in row.cells:
                _set_cell_shading(cell, color_hex)


# ──────────────────────────────────────────────────────────────────
# Main report writer
# ──────────────────────────────────────────────────────────────────

class DocxReportWriter:
    """Generates a comprehensive Word report for Part 23 certification
    loads analysis.

    Parameters
    ----------
    config : AircraftConfig
        Aircraft configuration data.
    vn_diagram : VnDiagram
        Computed V-n diagram.
    matrix : LoadCaseMatrix
        Load case matrix.
    batch_result : BatchResult
        Solver execution results.
    envelope_proc : EnvelopeProcessor
        Envelope processing results.
    report : CertificationReport
        High-level report data.
    model : BDFModel, optional
        Parsed BDF model for structural details.
    plot_dir : str, optional
        Directory containing generated plot PNGs.
    vmt_data : dict, optional
        VMT integration data.
    analysis_time : datetime, optional
        Timestamp of analysis execution.
    """

    def __init__(
        self,
        config: AircraftConfig,
        vn_diagram: VnDiagram,
        matrix: LoadCaseMatrix,
        batch_result: BatchResult,
        envelope_proc: EnvelopeProcessor,
        report: CertificationReport,
        model=None,
        plot_dir: str = None,
        vmt_data: dict = None,
        analysis_time: datetime = None,
        monitoring_stations: dict = None,
    ):
        if not HAS_DOCX:
            raise ImportError(
                "python-docx is required: pip install python-docx"
            )

        self.config = config
        self.vn = vn_diagram
        self.matrix = matrix
        self.batch = batch_result
        self.proc = envelope_proc
        self.report = report
        self.model = model
        self.plot_dir = plot_dir or "."
        self.vmt_data = vmt_data or {}
        self.analysis_time = analysis_time or datetime.now()
        self.monitoring_stations = monitoring_stations or {}

        self.doc = Document()
        self._setup_styles()

    # ──────────────────────────────────────────────────────────────
    # Document setup
    # ──────────────────────────────────────────────────────────────

    def _setup_styles(self):
        """Configure document default styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

        # Heading styles
        for level in range(1, 5):
            heading_style = self.doc.styles[f'Heading {level}']
            heading_style.font.name = 'Calibri'
            heading_style.font.color.rgb = RGBColor(0x1A, 0x5C, 0x97)

        # Page setup
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # ──────────────────────────────────────────────────────────────
    # Utility methods
    # ──────────────────────────────────────────────────────────────

    def _add_image(self, filename: str, width_inches: float = 6.0,
                   caption: str = None):
        """Add an image from plot_dir if it exists."""
        path = os.path.join(self.plot_dir, filename)
        if not os.path.isfile(path):
            # Try without directory prefix
            if os.path.isfile(filename):
                path = filename
            else:
                p = self.doc.add_paragraph(
                    f"[Figure not available: {filename}]"
                )
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                return

        self.doc.add_picture(path, width=Inches(width_inches))
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def _add_equation(self, equation_text: str, description: str = None):
        """Add a formatted equation block."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(equation_text)
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
        run.font.italic = True

        if description:
            d = self.doc.add_paragraph()
            d.alignment = WD_ALIGN_PARAGRAPH.LEFT
            d.paragraph_format.left_indent = Cm(2.0)
            run = d.add_run(f"where: {description}")
            run.font.size = Pt(9)
            run.font.italic = True

    def _add_bullet(self, text: str, level: int = 0, bold_prefix: str = None):
        """Add a bullet point."""
        p = self.doc.add_paragraph(style='List Bullet')
        if level > 0:
            p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
        if bold_prefix:
            run = p.add_run(bold_prefix)
            run.bold = True
            p.add_run(f" {text}")
        else:
            p.add_run(text)

    def _section_break(self):
        """Add a page break."""
        self.doc.add_page_break()

    # ──────────────────────────────────────────────────────────────
    # Report sections
    # ──────────────────────────────────────────────────────────────

    def _write_cover_page(self):
        """Title page."""
        # Spacing before title
        for _ in range(6):
            self.doc.add_paragraph()

        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CERTIFICATION LOADS ANALYSIS REPORT")
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x97)

        # Subtitle
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("FAA Part 23 — Normal Category Aircraft")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        self.doc.add_paragraph()

        # Aircraft name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("KC-100 Light Aircraft")
        run.font.size = Pt(18)
        run.font.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Info table
        info_table = self.doc.add_table(rows=5, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_data = [
            ("Document Type", "Certification Loads Analysis"),
            ("Certification Basis", "14 CFR Part 23"),
            ("Analysis Tool", "NastAero — Structural Analysis Framework"),
            ("Analysis Date",
             self.analysis_time.strftime("%Y-%m-%d %H:%M:%S")),
            ("Solver", "SOL 144 Static Aeroelastic Trim"),
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            for cell in info_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        self._section_break()

    def _write_toc_placeholder(self):
        """Table of contents placeholder."""
        self.doc.add_heading("TABLE OF CONTENTS", level=1)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "[Update this field in Microsoft Word: "
            "References → Table of Contents]"
        )
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.size = Pt(11)
        self._section_break()

    # ──── Chapter 1: Introduction ──────────────────────────────

    def _write_ch1_introduction(self):
        """Chapter 1: Introduction and scope."""
        self.doc.add_heading("1. INTRODUCTION", level=1)

        self.doc.add_heading("1.1 Purpose", level=2)
        self.doc.add_paragraph(
            "This report presents the results of the certification loads "
            "analysis for the KC-100 light aircraft, performed in accordance "
            "with the requirements of Title 14, Code of Federal Regulations "
            "(14 CFR), Part 23 — Airworthiness Standards: Normal Category "
            "Airplanes. The analysis determines the critical design loads for "
            "all primary structural components under the complete spectrum of "
            "flight, maneuver, gust, and ground loading conditions."
        )

        self.doc.add_heading("1.2 Scope of Analysis", level=2)
        self.doc.add_paragraph(
            "The analysis covers the following regulatory requirements:"
        )

        compliance = self.report.regulatory_compliance_matrix()
        covered = [e for e in compliance if e.status == "covered"]
        for entry in covered:
            self._add_bullet(
                f"{entry.title} — {entry.n_cases} load cases",
                bold_prefix=entry.section,
            )

        self.doc.add_paragraph()
        summary = self.report.summary()
        self.doc.add_paragraph(
            f"A total of {summary['total_cases']} load cases were generated "
            f"and analyzed, comprising {summary['flight_cases']} flight "
            f"conditions and {summary['landing_cases']} landing/ground "
            f"conditions. The solver convergence rate is "
            f"{summary['convergence_rate']*100:.1f}%, covering "
            f"{summary['far_sections_covered']} of "
            f"{summary['far_sections_total']} applicable FAR sections "
            f"({summary['compliance_rate']*100:.1f}% regulatory coverage)."
        )

        self.doc.add_heading("1.3 Certification Basis", level=2)
        self.doc.add_paragraph(
            "The certification loads analysis is conducted in accordance with "
            "the following regulatory standards:"
        )
        self._add_bullet(
            "Airworthiness Standards: Normal Category Airplanes",
            bold_prefix="14 CFR Part 23 —",
        )
        self._add_bullet(
            "Limit maneuvering load factors per §23.337",
            bold_prefix="Load Factors:",
        )
        self._add_bullet(
            "Pratt quasi-static gust formula per §23.341",
            bold_prefix="Gust Loads:",
        )
        self._add_bullet(
            "Rolling (§23.349) and Yaw (§23.351) maneuver conditions",
            bold_prefix="Lateral Loads:",
        )
        self._add_bullet(
            "Level, tail-down, one-wheel, side load per "
            "§23.471-§23.511",
            bold_prefix="Ground Loads:",
        )

        self.doc.add_heading("1.4 Analysis Tools", level=2)
        self.doc.add_paragraph(
            "The loads analysis was performed using the NastAero structural "
            "analysis framework, which implements the following capabilities:"
        )
        self._add_bullet(
            "Nastran-compatible BDF parser with full "
            "element/material/property support",
            bold_prefix="Structural Modeling:",
        )
        self._add_bullet(
            "Doublet-Lattice Method (DLM) for steady/unsteady "
            "aerodynamic computations",
            bold_prefix="Aerodynamic Analysis:",
        )
        self._add_bullet(
            "SOL 144 static aeroelastic trim solver with "
            "Two-G-matrix formulation and Schur complement",
            bold_prefix="Trim Solver:",
        )
        self._add_bullet(
            "Automated VMT (Shear-Moment-Torsion) integration "
            "along structural components",
            bold_prefix="Loads Integration:",
        )
        self._add_bullet(
            "6-DOF rigid-body flight dynamics simulation for "
            "dynamic maneuver and gust time-domain analysis",
            bold_prefix="Dynamic Simulation:",
        )

        # ---- 1.5 Analysis Procedure Flowchart ----
        self._write_analysis_procedure_flowchart()

        self._section_break()

    # ──── Chapter 1.5: Analysis Procedure Flowchart ────────────

    def _write_analysis_procedure_flowchart(self):
        """Section 1.5: loads analysis procedure as a block-diagram table."""
        self.doc.add_heading("1.5 Analysis Procedure Overview", level=2)
        self.doc.add_paragraph(
            "The following diagram summarises the end-to-end certification "
            "loads analysis procedure. Static and dynamic load cases are "
            "generated in parallel, merged into a single envelope, and "
            "processed through a common internal-loads pipeline."
        )

        # Build flowchart with tables (works in any docx viewer)
        # ---- Row 0: Inputs ----
        t = self.doc.add_table(rows=1, cols=3)
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        boxes_row0 = [
            ("① FE Model (BDF)\n"
             "─────────────\n"
             "• Structural elements\n"
             "• CONM2 mass dist.\n"
             "• Aero panels (CAERO)\n"
             "• Splines (SPLINE)"),
            ("② Aircraft Config\n"
             "─────────────\n"
             "• Weight / CG\n"
             "• Speed schedule\n"
             "  (VA, VB, VC, VD)\n"
             "• Control surface limits"),
            ("③ FAR 23 Requirements\n"
             "─────────────\n"
             "• §23.333 V-n diagram\n"
             "• §23.337 Load factors\n"
             "• §23.341 Gust loads\n"
             "• §23.349/351 Lateral"),
        ]
        for i, txt in enumerate(boxes_row0):
            t.rows[0].cells[i].text = txt
            _set_cell_shading(t.rows[0].cells[i], "D5E8D4")  # light green
            for p in t.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7.5)
        _set_table_style(t)

        # Arrow
        p_arrow = self.doc.add_paragraph()
        p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_arrow.add_run("▼")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x97)
        p_arrow.paragraph_format.space_before = Pt(2)
        p_arrow.paragraph_format.space_after = Pt(2)

        # ---- Row 1: V-n Diagram + VLM Derivatives ----
        t2 = self.doc.add_table(rows=1, cols=2)
        t2.style = 'Table Grid'
        t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        boxes_row1 = [
            ("④ V-n Diagram Construction\n"
             "─────────────────────\n"
             "• Maneuver envelope  (nz = ±(V/VS1)²)\n"
             "• Gust envelope  (Pratt formula, Kg alleviation)\n"
             "• Corner points:  A+, A−, D+, D−, G+, G−, VB±"),
            ("⑤ VLM Stability Derivatives\n"
             "─────────────────────\n"
             "• Perturbation:  CLα, Cmα, CLδe, Cmδe\n"
             "• Lateral VLM:  CYβ, Clβ, Cnβ, Clδa, CYδr, Cnδr\n"
             "• Empirical:  Cmq, Clp, Cnr  (damping terms)"),
        ]
        for i, txt in enumerate(boxes_row1):
            t2.rows[0].cells[i].text = txt
            _set_cell_shading(t2.rows[0].cells[i], "D6EAF8")  # light blue
            for p in t2.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7.5)
        _set_table_style(t2)

        # Split arrow
        p_arrow2 = self.doc.add_paragraph()
        p_arrow2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_arrow2.add_run("▼                                          ▼")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x97)
        p_arrow2.paragraph_format.space_before = Pt(2)
        p_arrow2.paragraph_format.space_after = Pt(2)

        # ---- Row 2: Static + Dynamic paths ----
        t3 = self.doc.add_table(rows=1, cols=2)
        t3.style = 'Table Grid'
        t3.alignment = WD_TABLE_ALIGNMENT.CENTER
        boxes_row2 = [
            ("⑥ Static Load Cases\n"
             "─────────────────────\n"
             "V-n corner points → trim conditions\n"
             "\n"
             "• Symmetric: nz at VA, VD\n"
             "• Gust:  ΔnG at VB, VC, VD\n"
             "• Rolling:  δa at VA, VC, VD  (§23.349)\n"
             "• Yaw:  δr at VA, VC, VD  (§23.351)\n"
             "• Landing:  gear loads  (§23.471-511)"),
            ("⑦ Dynamic 6-DOF Simulations\n"
             "─────────────────────\n"
             "Time-domain maneuver / gust profiles\n"
             "\n"
             "• Elevator pull-up & checked  (§23.331)\n"
             "• Abrupt roll  (§23.349)\n"
             "• Yaw maneuver  (§23.351)\n"
             "• Discrete 1-cos gust  (§23.341)\n"
             "→ extract critical time points\n"
             "  (max nz, |q̇|, |ṗ|, |ṙ|, |β|)"),
        ]
        for i, txt in enumerate(boxes_row2):
            t3.rows[0].cells[i].text = txt
            color = "FFF2CC" if i == 0 else "FCE4D6"  # yellow / orange
            _set_cell_shading(t3.rows[0].cells[i], color)
            for p in t3.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7.5)
        _set_table_style(t3)

        # Merge arrow
        p_arrow3 = self.doc.add_paragraph()
        p_arrow3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_arrow3.add_run("▼──────────── merge ────────────▼")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x97)
        p_arrow3.paragraph_format.space_before = Pt(2)
        p_arrow3.paragraph_format.space_after = Pt(2)

        # ---- Row 3: SOL 144 Solver ----
        t4 = self.doc.add_table(rows=1, cols=1)
        t4.style = 'Table Grid'
        t4.alignment = WD_TABLE_ALIGNMENT.CENTER
        t4.rows[0].cells[0].text = (
            "⑧ SOL 144 — Static Aeroelastic Trim Solver\n"
            "───────────────────────────────────────────\n"
            "All static + dynamic cases solved: "
            " [K − q·Q_aa] · {u} = {P} + q·[Q_ax]·{u_x}\n"
            "→ Elastic deformation → Nodal forces "
            "(aerodynamic + inertial + applied)"
        )
        _set_cell_shading(t4.rows[0].cells[0], "E8DAEF")  # light purple
        for p in t4.rows[0].cells[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(7.5)
        _set_table_style(t4)

        # Arrow
        p_arrow4 = self.doc.add_paragraph()
        p_arrow4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_arrow4.add_run("▼")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x97)
        p_arrow4.paragraph_format.space_before = Pt(2)
        p_arrow4.paragraph_format.space_after = Pt(2)

        # ---- Row 4: Post-processing ----
        t5 = self.doc.add_table(rows=1, cols=3)
        t5.style = 'Table Grid'
        t5.alignment = WD_TABLE_ALIGNMENT.CENTER
        boxes_row4 = [
            ("⑨ VMT Integration\n"
             "─────────────\n"
             "Shear, Bending\n"
             "Moment, Torsion\n"
             "at monitoring\n"
             "stations"),
            ("⑩ Envelope\n"
             "─────────────\n"
             "Max/Min over all\n"
             "cases per station\n"
             "→ Critical design\n"
             "loads & cases"),
            ("⑪ Output\n"
             "─────────────\n"
             "• Force cards (BDF)\n"
             "• Summary CSV\n"
             "• Certification\n"
             "  report (.docx)"),
        ]
        for i, txt in enumerate(boxes_row4):
            t5.rows[0].cells[i].text = txt
            _set_cell_shading(t5.rows[0].cells[i], "DAEEF3")  # light teal
            for p in t5.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7.5)
        _set_table_style(t5)

        self.doc.add_paragraph()
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run("Figure 1-1. Certification Loads Analysis Procedure")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ──── Chapter 2: Aircraft Description ──────────────────────

    def _write_ch2_aircraft(self):
        """Chapter 2: Aircraft description."""
        self.doc.add_heading("2. AIRCRAFT DESCRIPTION", level=1)

        self.doc.add_heading("2.1 General Configuration", level=2)

        cfg = self.config
        wc = cfg.weight_cg_conditions[0] if cfg.weight_cg_conditions else None

        # Aircraft parameters table
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0]
        hdr.cells[0].text = "Parameter"
        hdr.cells[1].text = "Value"
        hdr.cells[2].text = "Unit"
        _format_header_row(hdr)

        params = [
            ("Wing Area (S_ref)", f"{cfg.wing_area_m2:.1f}", "m²"),
            ("Mean Aerodynamic Chord (MAC)", f"{cfg.mean_chord_m:.2f}", "m"),
            ("Wing Span (estimated)", f"{cfg.wing_area_m2/cfg.mean_chord_m:.1f}", "m"),
            ("Aspect Ratio (estimated)", f"{(cfg.wing_area_m2/cfg.mean_chord_m**2):.1f}", "—"),
        ]

        if wc:
            params.extend([
                ("Maximum Takeoff Weight (MTOW)",
                 f"{wc.weight_N:.0f}", "N"),
                ("MTOW (mass)", f"{wc.weight_kg:.1f}", "kg"),
                ("MTOW (Imperial)", f"{wc.weight_lb:.0f}", "lb"),
                ("CG Position (x)", f"{wc.cg_x:.1f}", "mm"),
            ])

        _w = wc.weight_N if wc else 0.0
        params.extend([
            ("Wing Loading (W/S)",
             f"{cfg.wing_loading(_w) if _w else 0:.1f}", "N/m²"),
            ("CLα (VLM)", f"{cfg.CLalpha:.3f}", "/rad"),
            ("Positive Limit Load Factor (nz_max)",
             f"{cfg.nz_max(_w):.2f}" if _w else "—", "g"),
            ("Negative Limit Load Factor (nz_min)",
             f"{cfg.nz_min(_w):.2f}" if _w else "—", "g"),
        ])

        for label, value, unit in params:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[2].text = unit
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        _set_table_style(table)
        _add_alt_row_shading(table)

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Table 2-1. Aircraft General Parameters"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Design speeds
        self.doc.add_heading("2.2 Design Speed Schedule", level=2)
        self.doc.add_paragraph(
            "The design speed schedule defines the flight envelope boundaries "
            "in accordance with 14 CFR §23.333:"
        )

        spd = cfg.speeds
        spd_table = self.doc.add_table(rows=1, cols=4)
        spd_table.style = 'Table Grid'
        hdr = spd_table.rows[0]
        for i, h in enumerate(["Speed", "EAS (m/s)", "EAS (kt)", "Description"]):
            hdr.cells[i].text = h
        _format_header_row(hdr)

        speed_data = [
            ("VS1", spd.VS1, "Stall speed, clean configuration"),
            ("VA", spd.VA, "Design maneuvering speed"),
            ("VB", spd.VB, "Design speed for maximum gust intensity"),
            ("VC", spd.VC, "Design cruising speed"),
            ("VD", spd.VD, "Design dive speed"),
            ("VF", spd.VF, "Design flap speed"),
        ]
        for name, v, desc in speed_data:
            row = spd_table.add_row()
            row.cells[0].text = name
            row.cells[1].text = f"{v:.1f}" if v > 0 else "—"
            row.cells[2].text = f"{v*1.944:.1f}" if v > 0 else "—"
            row.cells[3].text = desc
            row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        _set_table_style(spd_table)
        _add_alt_row_shading(spd_table)

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Table 2-2. Design Speed Schedule"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Control surface limits
        self.doc.add_heading("2.3 Control Surface Deflection Limits", level=2)
        ctrl = cfg.ctrl_limits
        if ctrl:
            ctrl_table = self.doc.add_table(rows=1, cols=3)
            ctrl_table.style = 'Table Grid'
            hdr = ctrl_table.rows[0]
            for i, h in enumerate(
                    ["Control Surface", "Max Deflection (deg)", "Notes"]):
                hdr.cells[i].text = h
            _format_header_row(hdr)

            ctrl_data = [
                ("Aileron", f"±{ctrl.aileron_max_deg:.1f}",
                 "Full at VA; 2/3 at VC; 1/3 at VD"),
                ("Elevator", f"±{ctrl.elevator_max_deg:.1f}",
                 "Per trim requirement"),
                ("Rudder", f"±{ctrl.rudder_max_deg:.1f}",
                 "Full at VA; decreasing with speed"),
            ]
            for name, defl, notes in ctrl_data:
                row = ctrl_table.add_row()
                row.cells[0].text = name
                row.cells[1].text = defl
                row.cells[2].text = notes
                row.cells[1].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.CENTER

            _set_table_style(ctrl_table)
            _add_alt_row_shading(ctrl_table)

            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "Table 2-3. Control Surface Deflection Limits"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Landing gear
        self.doc.add_heading("2.4 Landing Gear Configuration", level=2)
        gear = cfg.landing_gear
        if gear:
            self.doc.add_paragraph(
                "The landing gear configuration parameters used for ground "
                "loads analysis are summarized below:"
            )
            gear_table = self.doc.add_table(rows=1, cols=3)
            gear_table.style = 'Table Grid'
            hdr = gear_table.rows[0]
            for i, h in enumerate(["Parameter", "Value", "Unit"]):
                hdr.cells[i].text = h
            _format_header_row(hdr)

            gear_data = [
                ("Main Gear Node IDs",
                 str(gear.main_gear_node_ids), "—"),
                ("Nose Gear Node IDs",
                 str(gear.nose_gear_node_ids), "—"),
                ("Main Gear Position (x)",
                 f"{gear.main_gear_x:.1f}", "mm"),
                ("Nose Gear Position (x)",
                 f"{gear.nose_gear_x:.1f}", "mm"),
                ("Strut Efficiency (η)",
                 f"{gear.strut_efficiency:.2f}", "—"),
                ("Stroke", f"{gear.stroke:.3f}", "m"),
                ("Design Sink Rate",
                 f"{gear.sink_rate_fps:.1f}", "ft/s"),
            ]
            for label, val, unit in gear_data:
                row = gear_table.add_row()
                row.cells[0].text = label
                row.cells[1].text = val
                row.cells[2].text = unit
                row.cells[1].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[2].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.CENTER

            _set_table_style(gear_table)
            _add_alt_row_shading(gear_table)

            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "Table 2-4. Landing Gear Configuration"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Structural model
        self.doc.add_heading("2.5 Structural Finite Element Model", level=2)
        if self.model:
            n_nodes = len(self.model.nodes)
            n_elem = len(self.model.elements)
            n_mass = len(self.model.masses)
            self.doc.add_paragraph(
                f"The aircraft structural model consists of {n_nodes:,} nodes, "
                f"{n_elem:,} elements, and {n_mass:,} concentrated mass "
                f"elements (CONM2). The finite element model is defined in "
                f"Nastran Bulk Data Format (BDF) and includes shell, beam, "
                f"and bar elements to represent the primary structure."
            )

            model_table = self.doc.add_table(rows=1, cols=2)
            model_table.style = 'Table Grid'
            hdr = model_table.rows[0]
            hdr.cells[0].text = "Model Component"
            hdr.cells[1].text = "Count"
            _format_header_row(hdr)

            model_data = [
                ("Grid Points (GRID)", f"{n_nodes:,}"),
                ("Elements", f"{n_elem:,}"),
                ("Concentrated Masses (CONM2)", f"{n_mass:,}"),
                ("Materials (MAT1/MAT2/MAT8)", f"{len(self.model.materials):,}"),
                ("Properties (PSHELL/PBAR/PBEAM)",
                 f"{len(self.model.properties):,}"),
            ]

            # Element type breakdown
            elem_types = {}
            for eid, elem in self.model.elements.items():
                etype = getattr(elem, 'type', type(elem).__name__)
                elem_types[etype] = elem_types.get(etype, 0) + 1

            for etype, count in sorted(elem_types.items(),
                                         key=lambda x: -x[1]):
                model_data.append((f"  {etype}", f"{count:,}"))

            for label, val in model_data:
                row = model_table.add_row()
                row.cells[0].text = label
                row.cells[1].text = val
                row.cells[1].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT

            _set_table_style(model_table)
            _add_alt_row_shading(model_table)

            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "Table 2-5. Structural Model Summary"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            self.doc.add_paragraph(
                "Structural model data was not provided for this report."
            )

        # Aerodynamic model
        self.doc.add_heading("2.6 Aerodynamic Model", level=2)
        self.doc.add_paragraph(
            "The aerodynamic model uses the Doublet-Lattice Method (DLM) for "
            "computing the steady and unsteady aerodynamic influence "
            "coefficient (AIC) matrix. The DLM provides panel-based "
            "aerodynamic forces which are interpolated to the structural grid "
            "via spline interpolation (IPS — Infinite Plate Spline or TPS — "
            "Thin Plate Spline)."
        )

        if self.model:
            # CAERO panels
            n_caero = len(getattr(self.model, 'caeros', {}))
            n_aesurf = len(getattr(self.model, 'aesurfs', {}))
            n_spline = len(getattr(self.model, 'splines', {}))

            aero_table = self.doc.add_table(rows=1, cols=2)
            aero_table.style = 'Table Grid'
            hdr = aero_table.rows[0]
            hdr.cells[0].text = "Aerodynamic Component"
            hdr.cells[1].text = "Count"
            _format_header_row(hdr)

            aero_data = [
                ("CAERO1 Panels", f"{n_caero}"),
                ("AESURF Control Surfaces", f"{n_aesurf}"),
                ("Spline Connections", f"{n_spline}"),
                ("CLα (computed from VLM AIC)",
                 f"{cfg.CLalpha:.3f} /rad"),
                ("Reference Area (S_ref)",
                 f"{cfg.wing_area_m2:.1f} m²"),
                ("Reference Chord (c_ref)",
                 f"{cfg.mean_chord_m:.2f} m"),
            ]

            for label, val in aero_data:
                row = aero_table.add_row()
                row.cells[0].text = label
                row.cells[1].text = val
                row.cells[1].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT

            _set_table_style(aero_table)
            _add_alt_row_shading(aero_table)

            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "Table 2-6. Aerodynamic Model Summary"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # VMT Integration
        self.doc.add_heading("2.7 Load Reference Axes and VMT Integration",
                              level=2)
        self.doc.add_paragraph(
            "Internal loads (Shear, Bending Moment, and Torsion — collectively "
            "VMT) are computed by integrating nodal forces along predefined "
            "load reference axes for each structural component. The load "
            "reference axis is defined by the structural grid topology and "
            "oriented along the primary load path of each component."
        )

        self.doc.add_paragraph(
            "The VMT integration procedure at each spanwise station s is:"
        )

        self._add_equation(
            "V(s) = Σ Fz_i   for all nodes i outboard of station s",
            "V(s) = Transverse shear force, Fz_i = nodal force component "
            "normal to reference axis",
        )
        self._add_equation(
            "M(s) = Σ Fz_i × (y_i - y_s)   for all nodes i outboard of s",
            "M(s) = Bending moment, y_i = node coordinate along span, "
            "y_s = station coordinate",
        )
        self._add_equation(
            "T(s) = Σ Mx_i + Σ Fz_i × (x_i - x_ref)",
            "T(s) = Torsion, Mx_i = nodal moment about reference axis, "
            "x_ref = shear center",
        )

        # Fuselage CG-based integration
        wc = (self.config.weight_cg_conditions[0]
              if self.config.weight_cg_conditions else None)
        if wc:
            self.doc.add_paragraph(
                "For the fuselage, the VMT integration uses a CG-based "
                "bi-directional approach. Instead of integrating from one end "
                "to the other, the fuselage is split at the aircraft center "
                f"of gravity (CG_x = {wc.cg_x:.1f} mm) and the integration "
                "proceeds as follows:"
            )
            self._add_bullet(
                "nodes from the nose toward the CG (forward integration)",
                bold_prefix="Forward of CG:",
            )
            self._add_bullet(
                "nodes from the tail toward the CG (aft integration)",
                bold_prefix="Aft of CG:",
            )
            self.doc.add_paragraph(
                "This produces a VMT distribution that peaks near the CG and "
                "decreases to zero at both the nose and tail, which correctly "
                "represents the fuselage internal loads and is the standard "
                "presentation for fuselage loads analysis."
            )

        # Components
        if self.vmt_data:
            first_case = next(iter(self.vmt_data.values()), {})
            comp_names = list(first_case.keys()) if first_case else []
            if comp_names:
                self.doc.add_paragraph(
                    "The following structural components are monitored:"
                )
                for comp in comp_names:
                    self._add_bullet(comp)

        self._section_break()

    # ──── Chapter 3: Theory & Procedures ──────────────────────

    def _write_ch3_theory(self):
        """Chapter 3: Theory and procedures."""
        self.doc.add_heading("3. THEORY AND PROCEDURES", level=1)

        # 3.1 Load factor
        self.doc.add_heading("3.1 Limit Maneuvering Load Factors (§23.337)",
                              level=2)
        self.doc.add_paragraph(
            "The positive limit maneuvering load factor for normal category "
            "aircraft is determined per 14 CFR §23.337(a):"
        )
        self._add_equation(
            "n_z,max = min(3.8,  2.1 + 24,000 / W)",
            "W = maximum takeoff weight in pounds (lb), "
            "n_z,max = positive limit load factor",
        )

        wc = (self.config.weight_cg_conditions[0]
              if self.config.weight_cg_conditions else None)
        _w = wc.weight_N if wc else 0.0
        if wc:
            self.doc.add_paragraph(
                f"For the KC-100 at MTOW = {wc.weight_lb:.0f} lb:"
            )
            nz_calc = 2.1 + 24000.0 / wc.weight_lb
            self._add_equation(
                f"n_z,max = min(3.8,  2.1 + 24,000 / {wc.weight_lb:.0f}) "
                f"= min(3.8,  {nz_calc:.2f}) = "
                f"{self.config.nz_max(_w):.2f}",
            )

        self.doc.add_paragraph(
            "The negative limit load factor is:"
        )
        self._add_equation(
            "n_z,min = -0.4 × n_z,max",
            f"n_z,min = -0.4 × {self.config.nz_max(_w):.2f} "
            f"= {self.config.nz_min(_w):.2f}",
        )

        # 3.2 V-n diagram
        self.doc.add_heading("3.2 V-n Diagram Construction (§23.333)",
                              level=2)
        self.doc.add_paragraph(
            "The V-n (velocity vs. load factor) diagram defines the flight "
            "envelope within which the aircraft must withstand all loading "
            "conditions. It combines the maneuver envelope and the gust "
            "envelope to identify the critical corner points."
        )

        self.doc.add_heading("3.2.1 Maneuver Envelope", level=3)
        self.doc.add_paragraph(
            "The maneuver envelope is bounded by the stall speed curve "
            "and the limit load factors:"
        )
        self._add_equation(
            "n_stall(V) = ± (V / V_S1)²",
            "V_S1 = stall speed in clean configuration, "
            "V = equivalent airspeed (EAS)",
        )
        self._add_equation(
            "V_A = V_S1 × √n_z,max",
            "V_A = design maneuvering speed",
        )

        self.doc.add_heading("3.2.2 Gust Loads — Pratt Formula (§23.341)",
                              level=3)
        self.doc.add_paragraph(
            "Gust-induced incremental load factors are computed using the "
            "Pratt quasi-static gust formula as specified in §23.341:"
        )
        self._add_equation(
            "Δn = (ρ₀ × V × CLα × K_g × U_de) / (2 × W/S)",
            "ρ₀ = sea-level air density (1.225 kg/m³), "
            "V = equivalent airspeed (m/s)",
        )
        self.doc.add_paragraph(
            "The gust alleviation factor K_g accounts for aircraft response:"
        )
        self._add_equation(
            "K_g = 0.88 × μ_g / (5.3 + μ_g)",
            "μ_g = aircraft mass ratio = 2W / (ρ₀ × c̄ × g × S × CLα)",
        )

        self.doc.add_paragraph(
            "Design gust velocities per §23.333(c):"
        )
        gust_table = self.doc.add_table(rows=1, cols=3)
        gust_table.style = 'Table Grid'
        hdr = gust_table.rows[0]
        for i, h in enumerate(["Speed", "U_de (ft/s)", "U_de (m/s)"]):
            hdr.cells[i].text = h
        _format_header_row(hdr)

        gust_data = [
            ("VB", "66.0", "20.12"),
            ("VC", f"{self.config.gust_Ude_VC_fps:.1f}",
             f"{self.config.gust_Ude_VC_fps * 0.3048:.2f}"),
            ("VD", f"{self.config.gust_Ude_VD_fps:.1f}",
             f"{self.config.gust_Ude_VD_fps * 0.3048:.2f}"),
        ]
        for spd, ude_fps, ude_ms in gust_data:
            row = gust_table.add_row()
            row.cells[0].text = spd
            row.cells[1].text = ude_fps
            row.cells[2].text = ude_ms
            for c in row.cells:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        _set_table_style(gust_table)
        _add_alt_row_shading(gust_table)
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Table 3-1. Design Gust Velocities"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # V-n diagram figure
        self._add_image("01_vn_diagram.png", width_inches=5.5,
                          caption="Figure 3-1. V-n Diagram "
                          "(Maneuver + Gust Envelope)")

        # Corner points table
        if self.vn and self.vn.corner_points:
            self.doc.add_paragraph()
            self.doc.add_heading("3.2.3 V-n Corner Points", level=3)

            cp_table = self.doc.add_table(rows=1, cols=4)
            cp_table.style = 'Table Grid'
            hdr = cp_table.rows[0]
            for i, h in enumerate(
                    ["Label", "V_EAS (m/s)", "nz (g)", "Category"]):
                hdr.cells[i].text = h
            _format_header_row(hdr)

            for pt in self.vn.corner_points:
                row = cp_table.add_row()
                row.cells[0].text = pt.label
                row.cells[1].text = f"{pt.V_eas:.1f}"
                row.cells[2].text = f"{pt.nz:+.3f}"
                row.cells[3].text = pt.category
                row.cells[1].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[2].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[3].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.CENTER

            _set_table_style(cp_table)
            _add_alt_row_shading(cp_table)
            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "Table 3-2. V-n Diagram Corner Points"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 3.3 Trim solver
        self.doc.add_heading("3.3 Static Aeroelastic Trim Analysis "
                              "(SOL 144)", level=2)
        self.doc.add_paragraph(
            "The trim solution determines the aircraft equilibrium state at "
            "each flight condition. The formulation uses the two-G-matrix "
            "approach with Schur complement for computational efficiency."
        )
        self.doc.add_paragraph(
            "The aeroelastic equilibrium equation is:"
        )
        self._add_equation(
            "[K_aa - q × Q_aa] × {u_a} = {P_a} + q × [Q_ax] × {u_x}",
            "K_aa = structural stiffness, Q_aa = AIC matrix, "
            "q = dynamic pressure, P_a = applied loads, "
            "u_x = rigid body + control parameters",
        )

        self.doc.add_paragraph(
            "The trim variables include angle of attack (ANGLEA), elevator "
            "deflection (ELEV), and optionally aileron (ARON) and rudder "
            "(RUD) for lateral load cases. The equilibrium constraints "
            "enforce force and moment balance at the specified load factor."
        )

        self.doc.add_paragraph(
            "The trim Right-Hand Side is scaled by the load factor:"
        )
        self._add_equation(
            "RHS_trim = n_z × W",
            "n_z = design load factor, W = aircraft weight",
        )

        # 3.4 Internal loads
        self.doc.add_heading("3.4 Internal Loads Computation", level=2)
        self.doc.add_paragraph(
            "After the trim solution, nodal forces are extracted by combining "
            "aerodynamic, inertial, and applied loads at each structural "
            "grid point:"
        )
        self._add_equation(
            "F_combined = F_aero + F_inertia + F_applied",
            "F_aero = aerodynamic forces (from DLM), "
            "F_inertia = n_z × m × g (mass × gravity × load factor), "
            "F_applied = external applied loads",
        )

        self.doc.add_paragraph(
            "The combined nodal forces are then integrated along each "
            "structural component to obtain the spanwise distribution of "
            "Shear (V), Bending Moment (M), and Torsion (T)."
        )

        # 3.5 Lateral loads
        self.doc.add_heading("3.5 Lateral Maneuver Loads", level=2)

        self.doc.add_heading("3.5.1 Rolling Conditions (§23.349)", level=3)
        self.doc.add_paragraph(
            "Rolling maneuver loads are computed using a conservative "
            "quasi-static approach. The trim solver computes the symmetric "
            "pitch equilibrium with a fixed aileron deflection, which "
            "produces asymmetric wing loads through the DLM aerodynamics. "
            "The roll rate alleviation effect is conservatively neglected."
        )
        self._add_bullet(
            f"Full deflection (±{self.config.ctrl_limits.aileron_max_deg:.0f}°)",
            bold_prefix="At VA:",
        )
        self._add_bullet("2/3 of maximum deflection", bold_prefix="At VC:")
        self._add_bullet("1/3 of maximum deflection", bold_prefix="At VD:")

        self.doc.add_heading("3.5.2 Yaw Conditions (§23.351)", level=3)
        self.doc.add_paragraph(
            "Yaw maneuver loads are similarly computed with fixed rudder "
            "deflection during symmetric pitch trim. The rudder aerodynamic "
            "forces generate vertical tail loads directly through the DLM."
        )
        self._add_bullet(
            f"Full deflection (±{self.config.ctrl_limits.rudder_max_deg:.0f}°)",
            bold_prefix="At VA:",
        )
        self._add_bullet("2/3 of maximum deflection", bold_prefix="At VC:")
        self._add_bullet("1/3 of maximum deflection", bold_prefix="At VD:")

        # 3.6 Ground loads
        self.doc.add_heading("3.6 Ground Loads (§23.471-§23.511)", level=2)
        self.doc.add_paragraph(
            "Landing and ground handling loads are computed using a "
            "quasi-static approach where gear reaction forces are applied as "
            "concentrated loads at the landing gear attachment nodes, combined "
            "with the inertial loads on the entire airframe."
        )
        self._add_equation(
            "n_z,land = W_descent / (2 × η × d × W / g)",
            "η = strut efficiency, d = stroke, "
            "W_descent = kinetic energy from sink rate",
        )

        self.doc.add_paragraph(
            "Landing conditions analyzed include:"
        )
        self._add_bullet("Level landing — 3-point contact", bold_prefix="§23.479:")
        self._add_bullet("Tail-down landing — main gear only", bold_prefix="§23.481:")
        self._add_bullet("One-wheel landing — single side, 0.75×nz", bold_prefix="§23.483:")
        self._add_bullet("Side load — lateral tire reaction", bold_prefix="§23.485:")
        self._add_bullet("Rebound — springback loads", bold_prefix="§23.487:")

        self._section_break()

    # ──── Chapter 4: Load Case Generation ─────────────────────

    def _write_ch4_load_cases(self):
        """Chapter 4: Load case matrix."""
        self.doc.add_heading("4. LOAD CASE GENERATION", level=1)

        self.doc.add_heading("4.1 Load Case Matrix Overview", level=2)
        summary = self.matrix.summary()
        total = self.matrix.total_cases
        n_flight = len(self.matrix.flight_cases)
        n_landing = len(self.matrix.landing_cases)

        self.doc.add_paragraph(
            f"A total of {total} load cases were generated for this "
            f"analysis. The load case matrix systematically covers all "
            f"V-n diagram corner points, combined with the aircraft weight/"
            f"CG conditions and analysis altitudes."
        )

        # Summary table
        summ_table = self.doc.add_table(rows=1, cols=3)
        summ_table.style = 'Table Grid'
        hdr = summ_table.rows[0]
        for i, h in enumerate(["Category", "Number of Cases", "FAR Section"]):
            hdr.cells[i].text = h
        _format_header_row(hdr)

        # Collect FAR sections per category
        cat_far = {}
        for c in self.matrix.flight_cases:
            cat = c.category
            far = c.far_section
            if cat not in cat_far:
                cat_far[cat] = set()
            if far:
                cat_far[cat].add(far)
        for c in self.matrix.landing_cases:
            far = c.far_section
            if "landing" not in cat_far:
                cat_far["landing"] = set()
            if far:
                cat_far["landing"].add(far)

        for cat, count in sorted(summary.items()):
            row = summ_table.add_row()
            row.cells[0].text = cat.capitalize()
            row.cells[1].text = str(count)
            fars = cat_far.get(cat, set())
            row.cells[2].text = ", ".join(sorted(fars)) if fars else "—"
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Total row
        row = summ_table.add_row()
        row.cells[0].text = "TOTAL"
        row.cells[1].text = str(total)
        row.cells[2].text = ""
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        _set_table_style(summ_table)
        _add_alt_row_shading(summ_table)
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Table 4-1. Load Case Matrix Summary"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Case matrix plots
        self._add_image("02_case_matrix_summary.png", width_inches=6.0,
                          caption="Figure 4-1. Load Case Distribution by "
                          "Category, Mach-nz, Altitude, and Counts")

        # 4.2 Weight/CG conditions
        self.doc.add_heading("4.2 Weight and CG Conditions", level=2)
        wc_table = self.doc.add_table(rows=1, cols=4)
        wc_table.style = 'Table Grid'
        hdr = wc_table.rows[0]
        for i, h in enumerate(
                ["Condition", "Weight (N)", "Weight (kg)", "CG_x (mm)"]):
            hdr.cells[i].text = h
        _format_header_row(hdr)

        for wc in self.config.weight_cg_conditions:
            row = wc_table.add_row()
            row.cells[0].text = wc.label
            row.cells[1].text = f"{wc.weight_N:,.0f}"
            row.cells[2].text = f"{wc.weight_kg:,.1f}"
            row.cells[3].text = f"{wc.cg_x:.1f}"
            for j in (1, 2, 3):
                row.cells[j].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT

        _set_table_style(wc_table)
        _add_alt_row_shading(wc_table)
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Table 4-2. Weight and CG Conditions"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 4.3 Analysis altitudes
        self.doc.add_heading("4.3 Analysis Altitudes", level=2)
        self.doc.add_paragraph(
            "Load cases were analyzed at the following altitudes:"
        )
        for alt in self.config.altitudes_m:
            self._add_bullet(f"{alt:.0f} m ({alt*3.281:.0f} ft)")

        # 4.4 Case generation process
        self.doc.add_heading("4.4 Case Generation Process", level=2)
        self.doc.add_paragraph(
            "The load case generation follows a systematic process:"
        )
        self._add_bullet(
            "Compute V-n diagram for each weight/CG condition and altitude "
            "combination to determine the flight envelope corner points."
        )
        self._add_bullet(
            "For each corner point (symmetric maneuver, gust, flap), "
            "generate a trim case with the corresponding load factor (nz) "
            "and equivalent airspeed (V_EAS)."
        )
        self._add_bullet(
            "Add rolling maneuver cases at VA, VC, VD with speed-dependent "
            "aileron deflection (full at VA, 2/3 at VC, 1/3 at VD)."
        )
        self._add_bullet(
            "Add yaw maneuver cases at VA, VC, VD with speed-dependent "
            "rudder deflection."
        )
        self._add_bullet(
            "Add checked maneuver cases (pull-up/push-over) per §23.331(c)."
        )
        self._add_bullet(
            "Generate landing and ground handling conditions per "
            "§23.471-§23.511."
        )

        self._section_break()

    # ──── Chapter 5: Solver Execution ─────────────────────────

    def _write_ch5_solver(self):
        """Chapter 5: Solver execution and results."""
        self.doc.add_heading("5. SOLVER EXECUTION", level=1)

        self.doc.add_heading("5.1 Execution Summary", level=2)
        n_total = self.batch.n_total
        n_conv = self.batch.n_converged
        wall = getattr(self.batch, 'wall_time_s', 0)

        self.doc.add_paragraph(
            f"The SOL 144 static aeroelastic trim solver was executed for "
            f"{n_total} load cases. The analysis completed in "
            f"{wall:.1f} seconds with {n_conv}/{n_total} cases converging "
            f"({100*n_conv/max(n_total,1):.1f}% convergence rate)."
        )

        # Convergence by category
        self.doc.add_heading("5.2 Convergence by Category", level=2)
        by_cat = {}
        for r in self.batch.case_results:
            cat = r.category
            if cat not in by_cat:
                by_cat[cat] = {"total": 0, "converged": 0}
            by_cat[cat]["total"] += 1
            if r.converged:
                by_cat[cat]["converged"] += 1

        conv_table = self.doc.add_table(rows=1, cols=4)
        conv_table.style = 'Table Grid'
        hdr = conv_table.rows[0]
        for i, h in enumerate(
                ["Category", "Total", "Converged", "Rate (%)"]):
            hdr.cells[i].text = h
        _format_header_row(hdr)

        for cat, info in sorted(by_cat.items()):
            row = conv_table.add_row()
            row.cells[0].text = cat.capitalize()
            row.cells[1].text = str(info["total"])
            row.cells[2].text = str(info["converged"])
            rate = 100 * info["converged"] / max(info["total"], 1)
            row.cells[3].text = f"{rate:.1f}"
            for j in (1, 2, 3):
                row.cells[j].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT

        _set_table_style(conv_table)
        _add_alt_row_shading(conv_table)
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Table 5-1. Solver Convergence by Category"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Force magnitude check
        self.doc.add_heading("5.3 Force Magnitude Verification", level=2)
        self.doc.add_paragraph(
            "As a sanity check, the maximum nodal force magnitude for each "
            "category is reported:"
        )

        force_table = self.doc.add_table(rows=1, cols=3)
        force_table.style = 'Table Grid'
        hdr = force_table.rows[0]
        for i, h in enumerate(["Category", "Max |F| (N)", "Status"]):
            hdr.cells[i].text = h
        _format_header_row(hdr)

        cat_maxf = {}
        for r in self.batch.case_results:
            if r.nodal_forces:
                max_f = max(np.linalg.norm(f[:3])
                            for f in r.nodal_forces.values())
                cat_maxf.setdefault(r.category, 0.0)
                cat_maxf[r.category] = max(cat_maxf[r.category], max_f)

        for cat, mf in sorted(cat_maxf.items()):
            row = force_table.add_row()
            row.cells[0].text = cat.capitalize()
            row.cells[1].text = f"{mf:,.0f}"
            status = "OK" if mf < 100000 else "HIGH"
            row.cells[2].text = status
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        _set_table_style(force_table)
        _add_alt_row_shading(force_table)
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Table 5-2. Maximum Nodal Force per Category"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._section_break()

    # ──── Chapter 6: VMT Results ──────────────────────────────

    def _write_ch6_vmt(self):
        """Chapter 6: VMT internal loads results."""
        self.doc.add_heading("6. VMT INTERNAL LOADS", level=1)

        self.doc.add_heading("6.1 VMT Overview", level=2)
        n_cases_vmt = len(self.vmt_data) if self.vmt_data else 0
        self.doc.add_paragraph(
            f"VMT (Shear-Moment-Torsion) internal loads were computed for "
            f"{n_cases_vmt} converged load cases by integrating the nodal "
            f"combined forces along each structural component."
        )

        if self.vmt_data:
            first_case = next(iter(self.vmt_data.values()), {})
            comp_names = list(first_case.keys()) if first_case else []

            # VMT max values table
            self.doc.add_heading("6.2 Maximum VMT Values", level=2)

            vmt_table = self.doc.add_table(rows=1, cols=4)
            vmt_table.style = 'Table Grid'
            hdr = vmt_table.rows[0]
            for i, h in enumerate(
                    ["Component", "Max |Shear| (N)",
                     "Max |Bending| (N·mm)", "Max |Torsion| (N·mm)"]):
                hdr.cells[i].text = h
            _format_header_row(hdr)

            for comp in comp_names:
                max_s = max_b = max_t = 0.0
                for cid, case_data in self.vmt_data.items():
                    if comp in case_data:
                        d = case_data[comp]
                        max_s = max(max_s, np.max(np.abs(d["shear"])))
                        max_b = max(max_b, np.max(np.abs(d["bending"])))
                        max_t = max(max_t, np.max(np.abs(d["torsion"])))
                row = vmt_table.add_row()
                row.cells[0].text = comp
                row.cells[1].text = f"{max_s:,.0f}"
                row.cells[2].text = f"{max_b:,.0f}"
                row.cells[3].text = f"{max_t:,.0f}"
                for j in (1, 2, 3):
                    row.cells[j].paragraphs[0].alignment = \
                        WD_ALIGN_PARAGRAPH.RIGHT

            _set_table_style(vmt_table)
            _add_alt_row_shading(vmt_table)
            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "Table 6-1. Maximum Absolute VMT Values per Component"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # VMT envelope plots
            self.doc.add_heading("6.3 VMT Envelope Diagrams", level=2)
            self.doc.add_paragraph(
                "The VMT envelope diagrams show the maximum and minimum "
                "shear, bending moment, and torsion distributions along "
                "each structural component, considering all analyzed load "
                "cases."
            )

            fig_num = 1
            for comp in comp_names:
                safe = comp.replace(' ', '_').lower()
                fname = f"03_vmt_envelope_{safe}.png"
                self._add_image(
                    fname, width_inches=5.5,
                    caption=f"Figure 6-{fig_num}. VMT Envelope — {comp}")
                fig_num += 1

            # Potato plots — multi-station
            self.doc.add_heading("6.4 Potato Plots (V-M Scatter)", level=2)
            self.doc.add_paragraph(
                "Potato plots show the Shear-Bending (V-M) envelope at "
                "critical monitoring stations along each structural "
                "component. Monitoring stations are automatically "
                "identified at control surface boundaries, landing gear "
                "locations, large mass items, wing/tail attachment points, "
                "and the aircraft CG. The convex hull defines the design "
                "envelope boundary, and extreme (critical) load cases are "
                "annotated."
            )

            # Monitoring station summary table
            if self.monitoring_stations:
                self.doc.add_heading(
                    "6.4.1 Monitoring Station Identification", level=3)
                sta_table = self.doc.add_table(rows=1, cols=4)
                for j, hdr in enumerate(
                        ["Component", "Station (mm)", "Label", "Reason"]):
                    sta_table.rows[0].cells[j].text = hdr
                _format_header_row(sta_table.rows[0])
                for comp_name in comp_names:
                    for ms in self.monitoring_stations.get(comp_name, []):
                        row = sta_table.add_row()
                        row.cells[0].text = comp_name
                        row.cells[1].text = f"{ms.position:.1f}"
                        row.cells[1].paragraphs[0].alignment = \
                            WD_ALIGN_PARAGRAPH.RIGHT
                        row.cells[2].text = ms.label
                        row.cells[3].text = ms.reason
                _set_table_style(sta_table)
                _add_alt_row_shading(sta_table)
                self.doc.add_paragraph()
                self.doc.add_paragraph(
                    "Table 6-2. Critical Monitoring Stations"
                ).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Per-component potato plots
            self.doc.add_heading(
                "6.4.2 Potato Plot Results", level=3)

            for comp in comp_names:
                safe = comp.replace(' ', '_').lower()
                comp_stations = self.monitoring_stations.get(comp, [])

                if comp_stations:
                    # Multi-station: images in potato_<comp>/ subdirectory
                    potato_subdir = f"potato_{safe}"
                    for idx, ms in enumerate(comp_stations):
                        sta_lbl = ms.label.replace(' ', '_').replace(
                            '/', '_')
                        fname = os.path.join(
                            potato_subdir,
                            f"04_potato_{safe}_{idx:02d}_{sta_lbl}.png")
                        self._add_image(
                            fname, width_inches=5.0,
                            caption=(
                                f"Figure 6-{fig_num}. Potato Plot — "
                                f"{comp}, {ms.label} "
                                f"(Station {ms.position:.0f} mm)"))
                        fig_num += 1
                else:
                    # Fallback: single mid-span plot (legacy)
                    fname = f"04_potato_{safe}.png"
                    self._add_image(
                        fname, width_inches=5.0,
                        caption=f"Figure 6-{fig_num}. Potato Plot — "
                        f"{comp} (Mid-Span)")
                    fig_num += 1

        self._section_break()

    # ──── Chapter 7: Envelope & Critical Cases ────────────────

    def _write_ch7_envelope(self):
        """Chapter 7: Envelope processing and critical case identification."""
        self.doc.add_heading(
            "7. CRITICAL DESIGN LOAD IDENTIFICATION", level=1
        )

        self.doc.add_heading("7.1 Envelope Processing Method", level=2)
        self.doc.add_paragraph(
            "The envelope processing identifies the critical design loads by "
            "examining the VMT distributions from all analyzed load cases. "
            "At each spanwise station of each structural component, the "
            "maximum and minimum values of shear (V), bending moment (M), "
            "and torsion (T) are extracted, along with the controlling load "
            "case information."
        )
        self.doc.add_paragraph(
            "The critical case identification follows this procedure:"
        )
        self._add_bullet(
            "For each structural component, extract V, M, T at all "
            "monitoring stations."
        )
        self._add_bullet(
            "At each station, record the load case producing the maximum "
            "and minimum values for each quantity (V, M, T)."
        )
        self._add_bullet(
            "Aggregate across all stations and components to identify the "
            "overall critical cases."
        )
        self._add_bullet(
            "Generate the convex hull envelope (potato plot) for combined "
            "V-M loading."
        )

        # Summary
        env_summary = self.proc.summary()
        self.doc.add_heading("7.2 Envelope Summary", level=2)
        self.doc.add_paragraph(
            f"The envelope processing identified "
            f"{env_summary.get('n_critical', 0)} critical conditions across "
            f"{len(env_summary.get('components', []))} structural components."
        )

        # Critical case frequency
        self.doc.add_heading("7.3 Critical Case Frequency", level=2)
        self.doc.add_paragraph(
            "The following chart shows which load cases most frequently "
            "appear as critical (envelope-defining) across all structural "
            "components and stations:"
        )
        self._add_image("05_critical_frequency.png", width_inches=5.5,
                          caption="Figure 7-1. Critical Case Frequency")

        # Category distribution
        cat_dist = self.proc.critical_category_distribution()
        if cat_dist:
            cat_dist_table = self.doc.add_table(rows=1, cols=3)
            cat_dist_table.style = 'Table Grid'
            hdr = cat_dist_table.rows[0]
            for i, h in enumerate(
                    ["Category", "Critical Count", "Percentage"]):
                hdr.cells[i].text = h
            _format_header_row(hdr)

            total_crit = sum(cat_dist.values())
            for cat, count in sorted(cat_dist.items(),
                                       key=lambda x: -x[1]):
                row = cat_dist_table.add_row()
                row.cells[0].text = cat.capitalize()
                row.cells[1].text = str(count)
                pct = 100.0 * count / max(total_crit, 1)
                row.cells[2].text = f"{pct:.1f}%"
                row.cells[1].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[2].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT

            _set_table_style(cat_dist_table)
            _add_alt_row_shading(cat_dist_table)
            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "Table 7-1. Critical Case Category Distribution"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Critical loads table
        self.doc.add_heading("7.4 Critical Design Loads Table", level=2)
        self.doc.add_paragraph(
            "The following table lists the critical design loads for each "
            "structural component. For each quantity (V, M, T) at each "
            "monitoring station, the maximum and minimum values are reported "
            "with the controlling load case."
        )

        crit_rows = self.report.critical_loads_table()
        if crit_rows:
            # Group by component
            from collections import defaultdict
            by_comp = defaultdict(list)
            for row in crit_rows:
                by_comp[row.component].append(row)

            table_num = 2
            for comp, rows in sorted(by_comp.items()):
                self.doc.add_heading(
                    f"7.4.{table_num-1} {comp}", level=3
                )

                # Select representative stations (max 10)
                all_stations = sorted(set(r.station for r in rows))
                if len(all_stations) > 10:
                    indices = np.linspace(0, len(all_stations) - 1, 10,
                                           dtype=int)
                    selected = [all_stations[i] for i in indices]
                else:
                    selected = all_stations

                ct = self.doc.add_table(rows=1, cols=8)
                ct.style = 'Table Grid'
                hdr = ct.rows[0]
                for i, h in enumerate([
                    "Station", "Qty",
                    "Max Value", "Max Case", "Max Cat.",
                    "Min Value", "Min Case", "Min Cat.",
                ]):
                    hdr.cells[i].text = h
                _format_header_row(hdr)

                for r in rows:
                    if r.station not in selected:
                        continue
                    row = ct.add_row()
                    row.cells[0].text = f"{r.station:.0f}"
                    row.cells[1].text = r.quantity
                    row.cells[2].text = f"{r.max_value:,.0f}"
                    row.cells[3].text = str(r.max_case_id)
                    row.cells[4].text = r.max_category[:6]
                    row.cells[5].text = f"{r.min_value:,.0f}"
                    row.cells[6].text = str(r.min_case_id)
                    row.cells[7].text = r.min_category[:6]
                    for j in (0, 2, 3, 5, 6):
                        row.cells[j].paragraphs[0].alignment = \
                            WD_ALIGN_PARAGRAPH.RIGHT
                    for j in (1, 4, 7):
                        row.cells[j].paragraphs[0].alignment = \
                            WD_ALIGN_PARAGRAPH.CENTER

                # Small font for compact tables
                for row in ct.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(7)

                _set_table_style(ct)
                _add_alt_row_shading(ct)
                self.doc.add_paragraph()
                self.doc.add_paragraph(
                    f"Table 7-{table_num}. Critical Loads — {comp}"
                ).alignment = WD_ALIGN_PARAGRAPH.CENTER
                table_num += 1

        # Critical cases detail
        all_critical = self.proc.get_critical_cases()
        if all_critical:
            self._write_critical_conditions_summary(all_critical)

        # ---- 7.6 Critical Flight Conditions ----
        self._write_critical_flight_conditions(all_critical)

        self._section_break()

    # ──── Chapter 7.5: Critical conditions summary (compact) ──

    def _write_critical_conditions_summary(self, all_critical):
        """Section 7.5: compact one-page summary of critical design loads.

        Shows a single table per component with max/min V, M, T,
        controlling case, and load category. The full station-by-station
        detail is deferred to Appendix B.
        """
        from collections import defaultdict

        self.doc.add_heading(
            "7.5 Critical Design Load Conditions Summary", level=2
        )
        self.doc.add_paragraph(
            "The following tables present the overall critical (envelope-"
            "defining) design loads for each structural component. Only the "
            "global extreme values per load quantity are shown; the full "
            "station-by-station breakdown is provided in Appendix B."
        )

        # Group by component
        by_comp = defaultdict(list)
        for cc in all_critical:
            by_comp[cc.component].append(cc)

        table_num = len(by_comp) + 2  # after 7-1 (category dist) and 7-2…

        for comp, cases in sorted(by_comp.items()):
            # Find the overall max/min for each quantity
            extremes = {}  # (quantity, extreme) -> CriticalCase
            for cc in cases:
                key = (cc.quantity, cc.extreme)
                if key not in extremes:
                    extremes[key] = cc
                else:
                    prev = extremes[key]
                    if cc.extreme == "max" and cc.value > prev.value:
                        extremes[key] = cc
                    elif cc.extreme == "min" and cc.value < prev.value:
                        extremes[key] = cc

            # Also count unique controlling case IDs for this component
            unique_case_ids = set(cc.case_id for cc in cases)
            n_stations = len(set(cc.station for cc in cases))

            self.doc.add_heading(comp, level=3)
            self.doc.add_paragraph(
                f"{len(cases)} critical conditions at {n_stations} monitoring "
                f"stations, driven by {len(unique_case_ids)} unique load cases."
            )

            # Summary table: Qty | Extreme | Value | Station | Case ID | Category | Label
            st = self.doc.add_table(rows=1, cols=7)
            st.style = 'Table Grid'
            hdr = st.rows[0]
            for i, h in enumerate([
                "Qty", "Extreme", "Value", "Station\n(mm)",
                "Case ID", "Category", "Label",
            ]):
                hdr.cells[i].text = h
            _format_header_row(hdr)

            for qty in ("V", "M", "T"):
                for ext in ("max", "min"):
                    cc = extremes.get((qty, ext))
                    if cc is None:
                        continue
                    cr = self.batch.get_result(cc.case_id)
                    label = (cr.label[:30] if cr and cr.label
                             else f"Case {cc.case_id}")
                    row = st.add_row()
                    row.cells[0].text = qty
                    row.cells[1].text = ext
                    row.cells[2].text = f"{cc.value:,.0f}"
                    row.cells[3].text = f"{cc.station:.0f}"
                    row.cells[4].text = str(cc.case_id)
                    cat_short = cc.category.replace("dynamic_", "dyn_")
                    row.cells[5].text = cat_short
                    row.cells[6].text = label

                    row.cells[0].paragraphs[0].alignment = \
                        WD_ALIGN_PARAGRAPH.CENTER
                    row.cells[1].paragraphs[0].alignment = \
                        WD_ALIGN_PARAGRAPH.CENTER
                    for j in (2, 3, 4):
                        row.cells[j].paragraphs[0].alignment = \
                            WD_ALIGN_PARAGRAPH.RIGHT

            for row in st.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(7.5)

            _set_table_style(st)
            _add_alt_row_shading(st)
            self.doc.add_paragraph()

        self.doc.add_paragraph(
            "Note: V = Shear (N), M = Bending Moment (N·mm), "
            "T = Torsion (N·mm). Complete station-by-station detail "
            "is in Appendix B."
        ).runs[0].font.italic = True

    # ──── Chapter 7.6: Critical flight conditions ────────────

    def _write_critical_flight_conditions(self, all_critical):
        """Section 7.6: flight state at each critical design condition.

        Shows altitude, velocity, load factors, angular rates,
        control-surface deflections, and AoA/sideslip for every
        unique critical load case.
        """
        if not all_critical:
            return

        # Collect unique case IDs that drive the envelope
        seen_ids = set()
        unique_cases = []
        for cc in all_critical:
            if cc.case_id in seen_ids:
                continue
            seen_ids.add(cc.case_id)
            cr = self.batch.get_result(cc.case_id)
            if cr is None:
                continue
            unique_cases.append((cc, cr))

        if not unique_cases:
            return

        # Switch to landscape for wide table
        new_section = self.doc.add_section(WD_ORIENT.LANDSCAPE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = Cm(29.7)
        new_section.page_height = Cm(21.0)
        new_section.left_margin = Cm(1.5)
        new_section.right_margin = Cm(1.5)
        new_section.top_margin = Cm(2.0)
        new_section.bottom_margin = Cm(2.0)

        self.doc.add_heading(
            "7.6 Critical Flight Conditions", level=2
        )
        self.doc.add_paragraph(
            "The following table documents the flight state at each "
            "critical design load condition. For cases derived from "
            "6-DOF dynamic simulations (case IDs ≥ 10000), the values "
            "correspond to the time instant at which the critical "
            "response peak occurred. For static trim cases, the values "
            "are the prescribed trim inputs."
        )

        # ---------- Table ----------
        headers = [
            "Case ID", "Category", "Label",
            "Alt\n(m)", "V_EAS\n(m/s)", "Mach",
            "nz\n(g)", "ny\n(g)",
            "α\n(°)", "β\n(°)",
            "p\n(°/s)", "q\n(°/s)", "r\n(°/s)",
            "ṗ\n(°/s²)", "q̇\n(°/s²)", "ṙ\n(°/s²)",
            "δe\n(°)", "δa\n(°)", "δr\n(°)",
        ]
        n_cols = len(headers)
        tbl = self.doc.add_table(rows=1, cols=n_cols)
        tbl.style = 'Table Grid'
        hdr_row = tbl.rows[0]
        for i, h in enumerate(headers):
            hdr_row.cells[i].text = h
        _format_header_row(hdr_row)

        for cc, cr in sorted(unique_cases, key=lambda x: x[1].case_id):
            fs = cr.flight_state or {}
            row = tbl.add_row()

            # Identifiers
            row.cells[0].text = str(cr.case_id)
            cat_short = cr.category.replace("dynamic_", "dyn_")
            row.cells[1].text = cat_short
            label_short = cr.label[:28] if cr.label else ""
            row.cells[2].text = label_short

            # Flight condition
            row.cells[3].text = f"{fs.get('altitude_m', cr.altitude_m):.0f}"
            row.cells[4].text = f"{fs.get('V_eas_m_s', 0):.1f}"
            row.cells[5].text = f"{fs.get('mach', cr.mach):.3f}"

            # Load factors
            row.cells[6].text = f"{fs.get('nz', cr.nz):+.2f}"
            row.cells[7].text = f"{fs.get('ny', 0):+.2f}"

            # Angles
            row.cells[8].text = f"{fs.get('alpha_deg', 0):.2f}"
            row.cells[9].text = f"{fs.get('beta_deg', 0):.2f}"

            # Angular rates (convert rad/s → deg/s for readability)
            p_dps = math.degrees(fs.get('p_rad_s', 0))
            q_dps = math.degrees(fs.get('q_rad_s', 0))
            r_dps = math.degrees(fs.get('r_rad_s', 0))
            row.cells[10].text = f"{p_dps:+.1f}"
            row.cells[11].text = f"{q_dps:+.1f}"
            row.cells[12].text = f"{r_dps:+.1f}"

            # Angular accelerations (convert rad/s² → deg/s²)
            pd_dps2 = math.degrees(fs.get('p_dot_rad_s2', 0))
            qd_dps2 = math.degrees(fs.get('q_dot_rad_s2', 0))
            rd_dps2 = math.degrees(fs.get('r_dot_rad_s2', 0))
            row.cells[13].text = f"{pd_dps2:+.1f}"
            row.cells[14].text = f"{qd_dps2:+.1f}"
            row.cells[15].text = f"{rd_dps2:+.1f}"

            # Control surface deflections
            row.cells[16].text = f"{fs.get('delta_e_deg', 0):+.1f}"
            row.cells[17].text = f"{fs.get('delta_a_deg', 0):+.1f}"
            row.cells[18].text = f"{fs.get('delta_r_deg', 0):+.1f}"

            # Right-align numeric cells
            for j in range(3, n_cols):
                row.cells[j].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT

        # Compact font (19-column table needs tight sizing)
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(6)

        _set_table_style(tbl)
        _add_alt_row_shading(tbl)
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Table 7-X. Flight State at Critical Design Load Conditions"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Dynamic-only detail: reason for criticality
        dyn_cases = [(cc, cr) for cc, cr in unique_cases
                     if cr.case_id >= 10000]
        if dyn_cases:
            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "For dynamic simulation cases, the critical response "
                "parameter that triggered each time point extraction:"
            )
            for cc, cr in sorted(dyn_cases,
                                  key=lambda x: x[1].case_id):
                fs = cr.flight_state or {}
                reason = fs.get('reason', '')
                mtype = fs.get('maneuver_type', cr.category)
                if reason:
                    self._add_bullet(
                        f" — {reason} (maneuver: {mtype})",
                        bold_prefix=f"Case {cr.case_id}",
                    )

        # Revert to portrait
        port_section = self.doc.add_section(WD_ORIENT.PORTRAIT)
        port_section.orientation = WD_ORIENT.PORTRAIT
        port_section.page_width = Cm(21.0)
        port_section.page_height = Cm(29.7)
        port_section.left_margin = Cm(2.0)
        port_section.right_margin = Cm(2.0)
        port_section.top_margin = Cm(2.0)
        port_section.bottom_margin = Cm(2.0)

    # ──── Chapter 8: Regulatory Compliance ────────────────────

    def _write_ch8_compliance(self):
        """Chapter 8: FAR Part 23 compliance matrix."""
        self.doc.add_heading("8. REGULATORY COMPLIANCE", level=1)

        self.doc.add_heading("8.1 FAR Part 23 Compliance Matrix", level=2)
        self.doc.add_paragraph(
            "The following table shows the regulatory compliance status for "
            "each applicable section of 14 CFR Part 23:"
        )

        compliance = self.report.regulatory_compliance_matrix()

        comp_table = self.doc.add_table(rows=1, cols=4)
        comp_table.style = 'Table Grid'
        hdr = comp_table.rows[0]
        for i, h in enumerate(
                ["FAR Section", "Title", "Cases", "Status"]):
            hdr.cells[i].text = h
        _format_header_row(hdr)

        status_colors = {
            "covered": "C8E6C9",      # green
            "partial": "FFF9C4",       # yellow
            "not_covered": "FFCDD2",   # red
        }

        for entry in compliance:
            row = comp_table.add_row()
            row.cells[0].text = entry.section
            row.cells[1].text = entry.title
            row.cells[2].text = str(entry.n_cases)
            row.cells[3].text = entry.status.replace("_", " ").title()
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Status cell coloring
            color = status_colors.get(entry.status)
            if color:
                _set_cell_shading(row.cells[3], color)

        _set_table_style(comp_table)
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Table 8-1. FAR Part 23 Compliance Matrix"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Compliance summary
        summary = self.report.summary()
        self.doc.add_heading("8.2 Compliance Summary", level=2)
        n_covered = summary['far_sections_covered']
        n_total = summary['far_sections_total']
        rate = summary['compliance_rate']
        self.doc.add_paragraph(
            f"Of the {n_total} applicable FAR sections, {n_covered} are "
            f"covered by the current analysis ({rate*100:.1f}% coverage). "
        )

        self._section_break()

    # ──── Chapter 9: Conclusions ──────────────────────────────

    def _write_ch9_conclusions(self):
        """Chapter 9: Conclusions."""
        self.doc.add_heading("9. CONCLUSIONS", level=1)

        summary = self.report.summary()
        self.doc.add_paragraph(
            "The FAA Part 23 certification loads analysis for the KC-100 "
            "light aircraft has been completed with the following results:"
        )

        self._add_bullet(
            f"{summary['total_cases']} load cases analyzed "
            f"({summary['flight_cases']} flight + "
            f"{summary['landing_cases']} landing/ground)",
        )
        self._add_bullet(
            f"Solver convergence rate: "
            f"{summary['convergence_rate']*100:.1f}%",
        )
        self._add_bullet(
            f"FAR sections covered: {summary['far_sections_covered']}/"
            f"{summary['far_sections_total']} "
            f"({summary['compliance_rate']*100:.1f}%)",
        )

        env = self.proc.summary()
        self._add_bullet(
            f"Critical design conditions identified: "
            f"{env.get('n_critical', 0)} across "
            f"{len(env.get('components', []))} structural components",
        )

        # Category insight
        cat_dist = self.proc.critical_category_distribution()
        if cat_dist:
            dominant = max(cat_dist, key=cat_dist.get)
            self._add_bullet(
                f"Dominant critical category: {dominant} "
                f"({cat_dist[dominant]} critical conditions)",
            )

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "The critical design loads identified in this analysis establish "
            "the load envelope for the structural substantiation of all "
            "primary aircraft components. These loads should be used as input "
            "for the stress analysis and structural sizing per the "
            "applicable airworthiness requirements."
        )

        self._section_break()

    # ──── Appendix A: Complete Load Case Listing ──────────────

    def _write_appendix_a(self):
        """Appendix A: Complete load case listing."""
        self.doc.add_heading("APPENDIX A — COMPLETE LOAD CASE LISTING",
                              level=1)

        # Flight cases
        self.doc.add_heading("A.1 Flight Load Cases", level=2)
        self.doc.add_paragraph(
            f"Total flight cases: {len(self.matrix.flight_cases)}"
        )

        if self.matrix.flight_cases:
            # Create table in batches to avoid very large single table
            ft = self.doc.add_table(rows=1, cols=7)
            ft.style = 'Table Grid'
            hdr = ft.rows[0]
            for i, h in enumerate([
                "Case ID", "Category", "FAR §",
                "V_EAS (m/s)", "Mach", "nz (g)", "Alt (m)",
            ]):
                hdr.cells[i].text = h
            _format_header_row(hdr)

            for c in self.matrix.flight_cases:
                tc = c.trim_condition
                row = ft.add_row()
                row.cells[0].text = str(c.case_id)
                row.cells[1].text = c.category
                row.cells[2].text = c.far_section or "—"
                if tc:
                    # Compute V_EAS from Mach if available
                    v_eas = ""
                    if hasattr(tc, 'V_eas'):
                        v_eas = f"{tc.V_eas:.1f}"
                    elif hasattr(tc, 'velocity'):
                        v_eas = f"{tc.velocity:.1f}"
                    row.cells[3].text = v_eas
                    row.cells[4].text = f"{tc.mach:.4f}"
                    row.cells[5].text = f"{tc.nz:+.3f}"
                else:
                    row.cells[3].text = "—"
                    row.cells[4].text = "—"
                    row.cells[5].text = "—"
                row.cells[6].text = f"{c.altitude_m:.0f}"

                # Right-align numeric columns
                for j in (0, 3, 4, 5, 6):
                    row.cells[j].paragraphs[0].alignment = \
                        WD_ALIGN_PARAGRAPH.RIGHT

            # Small font for appendix
            for row in ft.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(7)

            _set_table_style(ft)
            _add_alt_row_shading(ft)
            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "Table A-1. Flight Load Cases"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Landing cases
        self.doc.add_heading("A.2 Landing and Ground Load Cases", level=2)
        self.doc.add_paragraph(
            f"Total landing/ground cases: {len(self.matrix.landing_cases)}"
        )

        if self.matrix.landing_cases:
            lt = self.doc.add_table(rows=1, cols=5)
            lt.style = 'Table Grid'
            hdr = lt.rows[0]
            for i, h in enumerate([
                "Case ID", "Condition", "FAR §", "nz_CG", "Weight Cond.",
            ]):
                hdr.cells[i].text = h
            _format_header_row(hdr)

            for lc in self.matrix.landing_cases:
                row = lt.add_row()
                row.cells[0].text = str(lc.case_id)
                cond_type = getattr(lc, 'condition_type', None)
                if cond_type:
                    row.cells[1].text = (cond_type.value
                                          if hasattr(cond_type, 'value')
                                          else str(cond_type))
                else:
                    row.cells[1].text = getattr(lc, 'label', '—')
                row.cells[2].text = getattr(lc, 'far_section', '—') or '—'
                nz_cg = getattr(lc, 'nz_cg', None)
                row.cells[3].text = f"{nz_cg:.2f}" if nz_cg else "—"
                wc_obj = getattr(lc, 'weight_cg', None)
                row.cells[4].text = (wc_obj.label
                                      if wc_obj and hasattr(wc_obj, 'label')
                                      else "—")

                row.cells[0].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[3].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT

            for row in lt.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(7)

            _set_table_style(lt)
            _add_alt_row_shading(lt)
            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "Table A-2. Landing and Ground Load Cases"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Solver results summary per case
        self.doc.add_heading("A.3 Solver Results Summary", level=2)

        if self.batch.case_results:
            st = self.doc.add_table(rows=1, cols=6)
            st.style = 'Table Grid'
            hdr = st.rows[0]
            for i, h in enumerate([
                "Case ID", "Category", "nz", "Mach",
                "Converged", "Label",
            ]):
                hdr.cells[i].text = h
            _format_header_row(hdr)

            for cr in sorted(self.batch.case_results,
                               key=lambda r: r.case_id):
                row = st.add_row()
                row.cells[0].text = str(cr.case_id)
                row.cells[1].text = cr.category
                row.cells[2].text = f"{cr.nz:+.3f}"
                row.cells[3].text = f"{cr.mach:.4f}"
                row.cells[4].text = "Yes" if cr.converged else "No"
                row.cells[5].text = cr.label or "—"

                row.cells[0].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[2].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[3].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[4].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.CENTER

                # Highlight failed cases
                if not cr.converged:
                    _set_cell_shading(row.cells[4], "FFCDD2")

            for row in st.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(7)

            _set_table_style(st)
            _add_alt_row_shading(st)
            self.doc.add_paragraph()
            self.doc.add_paragraph(
                "Table A-3. Solver Results per Load Case"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ──── Appendix B: Full critical condition detail ──────────

    def _write_appendix_b(self):
        """Appendix B: Full critical design load conditions detail.

        Station-by-station listing of every envelope-defining load case
        for each structural component. This is the verbose version of
        Section 7.5 (which shows only the global extremes).
        """
        all_critical = self.proc.get_critical_cases()
        if not all_critical:
            return

        from collections import defaultdict

        self.doc.add_heading(
            "APPENDIX B — CRITICAL DESIGN LOAD CONDITIONS (DETAIL)",
            level=1,
        )
        self.doc.add_paragraph(
            "This appendix provides the complete station-by-station "
            "critical design load conditions for each structural "
            "component. At every monitoring station, the maximum and "
            "minimum values of shear (V), bending moment (M), and "
            "torsion (T) are listed with the controlling load case."
        )

        by_comp = defaultdict(list)
        for cc in all_critical:
            by_comp[cc.component].append(cc)

        table_num = 1
        for comp, cases in sorted(by_comp.items()):
            self.doc.add_heading(f"B.{table_num} {comp}", level=2)

            bt = self.doc.add_table(rows=1, cols=8)
            bt.style = 'Table Grid'
            hdr = bt.rows[0]
            for i, h in enumerate([
                "Station\n(mm)", "Qty", "Extreme",
                "Value", "Case ID", "Category",
                "nz\n(g)", "Label",
            ]):
                hdr.cells[i].text = h
            _format_header_row(hdr)

            # Sort by station, then quantity, then extreme
            qty_order = {"V": 0, "M": 1, "T": 2}
            ext_order = {"max": 0, "min": 1}
            sorted_cases = sorted(
                cases,
                key=lambda c: (
                    c.station,
                    qty_order.get(c.quantity, 9),
                    ext_order.get(c.extreme, 9),
                ),
            )

            for cc in sorted_cases:
                cr = self.batch.get_result(cc.case_id)
                label = (cr.label[:26] if cr and cr.label
                         else f"Case {cc.case_id}")
                nz_str = f"{cr.nz:+.2f}" if cr else "—"
                cat_short = cc.category.replace("dynamic_", "dyn_")

                row = bt.add_row()
                row.cells[0].text = f"{cc.station:.0f}"
                row.cells[1].text = cc.quantity
                row.cells[2].text = cc.extreme
                row.cells[3].text = f"{cc.value:,.0f}"
                row.cells[4].text = str(cc.case_id)
                row.cells[5].text = cat_short
                row.cells[6].text = nz_str
                row.cells[7].text = label

                for j in (0, 3, 4, 6):
                    row.cells[j].paragraphs[0].alignment = \
                        WD_ALIGN_PARAGRAPH.RIGHT
                for j in (1, 2, 5):
                    row.cells[j].paragraphs[0].alignment = \
                        WD_ALIGN_PARAGRAPH.CENTER

            for row in bt.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(6.5)

            _set_table_style(bt)
            _add_alt_row_shading(bt)
            self.doc.add_paragraph()
            self.doc.add_paragraph(
                f"Table B-{table_num}. Critical Design Loads — "
                f"{comp} (All Stations)"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER
            table_num += 1

    # ──────────────────────────────────────────────────────────────
    # Main generate method
    # ──────────────────────────────────────────────────────────────

    def generate(self, output_path: str = None) -> str:
        """Generate the complete Word report.

        Parameters
        ----------
        output_path : str, optional
            Output .docx file path. Auto-generated if None.

        Returns
        -------
        str
            Path to the generated report file.
        """
        if output_path is None:
            ts = self.analysis_time.strftime("%Y%m%d_%H%M%S")
            output_path = f"KC100_Cert_Report_{ts}.docx"

        # Write all sections
        self._write_cover_page()
        self._write_toc_placeholder()
        self._write_ch1_introduction()
        self._write_ch2_aircraft()
        self._write_ch3_theory()
        self._write_ch4_load_cases()
        self._write_ch5_solver()
        self._write_ch6_vmt()
        self._write_ch7_envelope()
        self._write_ch8_compliance()
        self._write_ch9_conclusions()
        self._write_appendix_a()
        self._write_appendix_b()

        # Footer: page numbers
        self._add_page_numbers()

        # Save
        self.doc.save(output_path)
        return output_path

    def _add_page_numbers(self):
        """Add page numbers to footer."""
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add page number field
            run = p.add_run()
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            fld_char_begin = parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
            )
            run._element.append(fld_char_begin)

            run2 = p.add_run()
            instr = parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve">'
                f' PAGE </w:instrText>'
            )
            run2._element.append(instr)

            run3 = p.add_run()
            fld_char_end = parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
            )
            run3._element.append(fld_char_end)


# ──────────────────────────────────────────────────────────────────
# Convenience function
# ──────────────────────────────────────────────────────────────────

def generate_cert_report(
    config: AircraftConfig,
    vn_diagram: VnDiagram,
    matrix: LoadCaseMatrix,
    batch_result: BatchResult,
    envelope_proc: EnvelopeProcessor,
    report: CertificationReport,
    model=None,
    plot_dir: str = None,
    vmt_data: dict = None,
    output_path: str = None,
    analysis_time: datetime = None,
    monitoring_stations: dict = None,
) -> str:
    """Generate a Word report for Part 23 certification loads analysis.

    Convenience wrapper around DocxReportWriter.

    Parameters
    ----------
    config : AircraftConfig
    vn_diagram : VnDiagram
    matrix : LoadCaseMatrix
    batch_result : BatchResult
    envelope_proc : EnvelopeProcessor
    report : CertificationReport
    model : BDFModel, optional
    plot_dir : str, optional
    vmt_data : dict, optional
    output_path : str, optional
    analysis_time : datetime, optional
    monitoring_stations : dict, optional

    Returns
    -------
    str
        Path to generated .docx file.
    """
    writer = DocxReportWriter(
        config=config,
        vn_diagram=vn_diagram,
        matrix=matrix,
        batch_result=batch_result,
        envelope_proc=envelope_proc,
        report=report,
        model=model,
        plot_dir=plot_dir,
        vmt_data=vmt_data,
        analysis_time=analysis_time,
        monitoring_stations=monitoring_stations,
    )
    return writer.generate(output_path)
